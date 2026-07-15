from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE = Path(__file__).resolve().parents[1]
ASSETS = BASE / "report_assets"
OUTPUT = BASE / "实验4传输协议与套接字应用编程_完成版.docx"


def set_run_font(run, size=12, bold=False, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def set_paragraph(paragraph, size=12, bold=False, align=None, first_line=True):
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(4)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(24)
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def add_p(doc, text="", size=12, bold=False, align=None, first_line=True):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    set_paragraph(paragraph, size=size, bold=bold, align=align, first_line=first_line)
    return paragraph


def add_heading_cn(doc, text, level=1):
    size = {1: 14, 2: 13, 3: 12}.get(level, 12)
    paragraph = add_p(doc, text, size=size, bold=True, first_line=False)
    paragraph.paragraph_format.space_before = Pt(8)
    return paragraph


def add_code(doc, title, path, max_chars=3600):
    add_heading_cn(doc, title, 3)
    text = Path(path).read_text(encoding="utf-8")
    if len(text) > max_chars:
        text = text[:max_chars].rstrip() + "\n// ……后续完整源码见项目对应文件。"
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(8.5)


def set_cell(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], str(value), False)
    doc.add_paragraph()
    return table


def add_image(doc, title, image_name, width=6.0):
    add_p(doc, title, size=10.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.add_run().add_picture(str(ASSETS / image_name), width=Inches(width))
    add_p(doc, f"图：{title}", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    add_p(doc, "深 圳 大 学 实 验 报 告", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    doc.add_paragraph()
    cover_rows = [
        ("课程名称：", "互联网编程"),
        ("实验项目名称：", "传输协议与套接字应用编程"),
        ("学院：", "计算机与软件学院"),
        ("专业：", "计算机科学与技术"),
        ("指导教师：", "李梦柯"),
        ("报告人 / 学号 / 班级：", "姚泽棉 / 2024150026 / 高性能班"),
        ("实验时间：", "2026/06/04——2026/06/25"),
        ("实验报告提交时间：", "2026/06/10"),
    ]
    for key, value in cover_rows:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.8
        run = paragraph.add_run(key)
        set_run_font(run, 14, True)
        run = paragraph.add_run(value)
        set_run_font(run, 14)
    add_p(doc, "教务处制", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    doc.add_page_break()

    add_heading_cn(doc, "一、实验目的与内容", 1)
    add_heading_cn(doc, "1. 实验目的", 2)
    for item in [
        "掌握服务器端 Socket 编程开发能力，理解 TCP 连接建立、请求读取和响应返回过程。",
        "掌握 HTTP 协议基本格式，能够根据请求方法、路径、请求头和请求体生成正确响应。",
        "掌握 Cookie 编程技术，能够在 Web 应用中保存并传递客户端会话状态。",
    ]:
        add_p(doc, item)

    add_heading_cn(doc, "2. 实验内容要求", 2)
    for item in [
        "基于 Java 编写 HTTP 服务器程序和 HTTP 客户端程序。",
        "服务器采用线程池处理客户端连接，支持多个客户端同时访问。",
        "实现 GET、HEAD、POST 三类请求，并对不同请求返回正确响应。",
        "在服务器上部署静态网站，能够返回 HTML、CSS、JS 和图片等多种类型资源。",
        "HTTP 客户端能够连接服务器并展示不同请求的响应结果。",
        "使用 Cookie 保存并传递会话状态，例如保存用户名、访问次数和最后访问时间。",
        "对服务器进行压力测试，分析可支持的并发访问能力和文件传输情况。",
    ]:
        add_p(doc, item)

    add_heading_cn(doc, "二、实验过程和代码与结果", 1)
    add_heading_cn(doc, "1. 设计思路与实验过程", 2)
    add_p(doc, "本实验采用 Java 原生 ServerSocket 实现 HTTP 服务器。服务器启动后监听 8080 端口，主线程只负责接收客户端连接，具体请求处理交给固定大小线程池完成。这样可以避免单个客户端阻塞整体服务，同时满足多客户端并发访问要求。")
    add_p(doc, "请求处理流程为：读取请求行和请求头；根据 Content-Length 读取 POST 请求体；解析请求方法、路径和 Cookie；路由到静态资源、登录接口或会话接口；最后拼接 HTTP/1.1 响应头和响应体返回客户端。")
    add_p(doc, "静态网站放在项目 www 目录下，包括 index.html、style.css、app.js 和 images/network.svg。服务器根据文件后缀设置 Content-Type，因此浏览器能够正确显示网页文本、样式、脚本和图像。")
    add_p(doc, "Cookie 会话使用 LAB4_SESSION 作为会话 ID。首次访问时服务器创建 UUID 并写入 Set-Cookie；后续请求携带该 Cookie 后，服务器可取回同一 Session，保存用户名、访问次数和最后访问时间。")

    add_heading_cn(doc, "2. 项目结构", 2)
    add_table(doc, ["路径", "作用"], [
        ("src/lab4/SimpleHttpServer.java", "HTTP 服务器端程序，负责监听端口、解析请求、返回静态资源、处理 Cookie 会话。"),
        ("src/lab4/HttpClientDemo.java", "HTTP 客户端程序，依次发送 GET、HEAD、POST 和会话查询请求。"),
        ("src/lab4/StressTester.java", "压力测试程序，用多线程模拟多个客户端并发访问服务器资源。"),
        ("www/index.html", "静态网站首页，包含表单和会话展示区域。"),
        ("www/style.css、www/app.js、www/images/network.svg", "静态样式、脚本和图像资源，用于验证多类型响应。"),
    ])

    add_heading_cn(doc, "3. 关键源码", 2)
    add_p(doc, "完整源码已放在项目 src/lab4 目录中。以下给出主要程序的关键代码片段。")
    add_code(doc, "服务器端核心代码 SimpleHttpServer.java", BASE / "src/lab4/SimpleHttpServer.java", 5200)
    add_code(doc, "客户端核心代码 HttpClientDemo.java", BASE / "src/lab4/HttpClientDemo.java", 3300)
    add_code(doc, "压力测试核心代码 StressTester.java", BASE / "src/lab4/StressTester.java", 3000)

    add_heading_cn(doc, "4. 运行结果与测试", 2)
    add_p(doc, "编译命令为 javac -encoding UTF-8 -d out src/lab4/*.java，运行服务器命令为 java -cp out lab4.SimpleHttpServer。服务器启动后监听 http://localhost:8080。")
    add_image(doc, "浏览器访问服务器首页，HTML、CSS、JS 正常加载", "browser-home.png", 6.2)
    add_p(doc, "访问 /images/network.svg 可以看到服务器返回的图像资源，响应类型为 image/svg+xml，说明服务器能够返回文本之外的图片类型文件。")
    add_image(doc, "浏览器访问 SVG 图片资源", "image-resource.png", 5.4)
    add_p(doc, "客户端程序依次发送 GET /、HEAD /images/network.svg、POST /login、GET /api/session 请求。运行结果显示：GET 能获得 HTML 页面；HEAD 返回状态码和响应头但响应体长度为 0；POST 能保存用户名；最后一次 GET 能通过 Cookie 读取同一会话中的用户名和访问次数。")
    add_image(doc, "HTTP 客户端 GET、HEAD、POST 和 Cookie 会话测试输出", "client-output.png", 6.2)

    add_heading_cn(doc, "5. 功能测试汇总", 2)
    add_table(doc, ["测试项", "请求示例", "预期结果", "实际结果"], [
        ("GET 请求", "GET /", "返回首页 HTML 文本", "状态码 200，返回 index.html 内容"),
        ("HEAD 请求", "HEAD /images/network.svg", "只返回响应头，不返回响应体", "状态码 200，Body bytes 为 0"),
        ("POST 请求", "POST /login", "接收表单并保存用户名", "返回“POST 请求处理成功”，用户名保存为 YaoZemian"),
        ("静态图片", "GET /images/network.svg", "返回图片资源", "浏览器可正常显示 SVG 图片"),
        ("Cookie 会话", "GET /api/session", "读取同一会话状态", "返回 sessionId、username、visitCount、lastVisit"),
        ("多客户端访问", "StressTester", "并发请求可被线程池处理", "50 并发时 500 次请求全部成功"),
    ])

    add_heading_cn(doc, "6. 服务器性能分析", 2)
    add_p(doc, "压力测试使用 StressTester 程序完成。测试方法为：启动服务器后，分别模拟 50、100、200 个并发客户端，每个客户端连续发送 10 次 GET 请求，请求路径在首页、CSS、JS、SVG 图片和 JSON 接口之间轮换，以同时覆盖文本文件、脚本文件、样式文件和图像文件传输。")
    add_table(doc, ["并发客户端数", "每客户端请求数", "总请求数", "成功数", "失败数", "耗时", "吞吐量"], [
        ("50", "10", "500", "500", "0", "2.978 s", "167.90 requests/s"),
        ("100", "10", "1000", "962", "38", "5.512 s", "181.42 requests/s"),
        ("200", "10", "2000", "1924", "76", "10.716 s", "186.64 requests/s"),
    ])
    add_image(doc, "压力测试运行结果", "stress-summary.png", 6.2)
    add_p(doc, "从测试结果看，服务器在 50 个并发客户端、500 次总请求下全部成功，说明固定线程池能够稳定支持该规模的同时访问。并发提升到 100 和 200 时，请求总吞吐量提高到约 181.42 和 186.64 requests/s，但开始出现少量失败。失败主要可能来自短时间内大量 Socket 连接同时建立，服务器线程池大小为 32，连接排队和系统端口资源会带来压力。")
    add_p(doc, "文件同时传输方面，压力测试路径同时包含 HTML、CSS、JS、SVG 和 JSON 五类资源。50 并发时全部成功，可认为服务器能够稳定处理至少 50 个客户端同时访问并传输多种静态文件；更高并发下仍能完成大部分请求，但为了提高可靠性，可以增大线程池、设置连接队列长度、加入连接超时和更高效的文件缓存。")

    add_heading_cn(doc, "三、实验总结", 1)
    add_p(doc, "本实验从零实现了一个简易 HTTP 服务器和客户端，对 HTTP 协议的请求行、请求头、响应头、响应体以及状态码有了更直接的理解。相比直接使用 Web 框架，使用 Socket 编程需要手动处理字节流、Content-Length、文件类型判断和连接关闭等细节，因此更能体现 HTTP 协议在传输层之上的工作方式。")
    add_p(doc, "实验中遇到的主要问题包括：一是 POST 请求体不能只读取请求头，需要根据 Content-Length 再读取指定长度的请求体；二是 HEAD 请求虽然不返回响应体，但 Content-Length 仍应保持为目标资源长度；三是静态文件路径必须进行 normalize 检查，避免客户端通过 ../ 访问网站根目录之外的文件；四是 Cookie 的读取和 Set-Cookie 的返回需要保持名称一致，才能在多次请求间维持会话状态。")
    add_p(doc, "通过本实验，我掌握了 Java 服务器端 Socket 编程、HTTP 三种请求方法处理、静态网站部署、Cookie 会话状态管理和压力测试方法。后续如果继续改进，可以加入持久连接、更多 MIME 类型、访问日志文件、缓存机制和更完善的异常处理，使服务器更接近真实 Web 服务器。")

    add_heading_cn(doc, "AI 使用说明", 1)
    add_p(doc, "本人在完成本实验报告过程中，使用了大语言模型作为辅助工具。")
    add_p(doc, "使用工具名称：ChatGPT / Codex。")
    add_p(doc, "使用用途：辅助理解实验要求、生成和检查 Java Socket HTTP 服务器与客户端代码、整理测试结果、辅助撰写实验报告。代码和报告内容已结合本地运行结果进行检查。")

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name is None:
                set_run_font(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.name is None:
                            set_run_font(run, size=10.5)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()

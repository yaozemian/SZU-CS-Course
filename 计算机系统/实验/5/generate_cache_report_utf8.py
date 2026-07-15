# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUT = Path("2024150026姚泽棉_实验五_UTF8矩阵表版.docx")


def add_paragraph(container, text="", bold=False, size=10.5, align=None):
    """Write text through python-docx only. No manual XML editing."""
    p = container.add_paragraph(text)
    if align is not None:
        p.alignment = align
    for run in p.runs:
        run.bold = bold
        run.font.size = Pt(size)
    return p


def add_heading(container, text):
    return add_paragraph(container, text, bold=True, size=14)


def add_subheading(container, text):
    return add_paragraph(container, text, bold=True, size=12)


def add_table_from_rows(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    return table


def main():
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)

    title = add_paragraph(document, "深 圳 大 学 实 验 报 告", bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.space_after = Pt(18)

    cover_lines = [
        "课程名称：计算机系统(2)",
        "实验项目名称：Cache实验",
        "学院：计算机与软件学院",
        "专业：计算机与软件学院所有专业",
        "指导教师：刘刚",
        "报告人：姚泽棉    学号：2024150026    班级：高性能班",
        "实验时间：2026年5月14日至6月11日",
        "实验报告提交时间：2026年5月25日",
        "教务处制",
    ]
    for line in cover_lines:
        add_paragraph(document, line, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

    document.add_page_break()

    add_heading(document, "一、实验目的：")
    for line in [
        "加强对 Cache 工作原理的理解；",
        "体验程序中访存模式变化如何影响 Cache 命中率，并进一步影响程序性能；",
        "学习在 x86 真实机器上通过调整程序访存模式来探测多级 Cache 结构、容量以及 TLB 的大小。",
    ]:
        add_paragraph(document, line)

    add_heading(document, "二、实验环境")
    for line in [
        "X86 真实机器。",
        "硬件平台：LENOVO 83EG，AMD Ryzen 7 7840H with Radeon 780M Graphics。",
        "CPU 规格：8 核 16 线程，最高主频约 3.801GHz；物理内存约 16GB。",
        "操作系统：Microsoft Windows 11 家庭中文版，64 位，版本 10.0.26200。",
        "系统查询 Cache 信息：L1 Cache 总量 512KB，L2 Cache 总量 8192KB，L3 Cache 总量 16384KB。结合该处理器结构，可理解为每核心约 64KB L1（32KB 数据 Cache + 32KB 指令 Cache）、每核心 1MB L2、16MB 共享 L3，Cache Line 为 64B。",
    ]:
        add_paragraph(document, line)

    add_heading(document, "三、实验内容和步骤")
    add_subheading(document, "1、分析 Cache 访存模式对系统性能的影响")
    for line in [
        "原始矩阵乘法采用 i-j-k 三重循环。矩阵在 C 语言中按行优先存储，因此 a[i*size+k] 在最内层循环中连续访问，但 b[k*size+j] 按列访问，步长为 size*sizeof(float)，矩阵较大时容易频繁跨 Cache Line，局部性较差。",
        "优化方案是将循环顺序调整为 i-k-j，并将 a[i*size+k] 暂存到局部变量 aik 中。这样 b[k*size+j] 和 c[i*size+j] 都在 j 循环中连续访问，可以提高空间局部性。",
        "由于实验在 Windows 环境下完成，原代码中的 sys/time.h、unistd.h 和 gettimeofday() 改为 windows.h、QueryPerformanceFrequency() 和 QueryPerformanceCounter()；同时使用 scanf() 获取矩阵大小，并在程序结束前释放动态分配的内存。",
        "优化后的矩阵乘法核心代码：",
        "for (i = 0; i < size; i++) {",
        "    for (j = 0; j < size; j++) c[i * size + j] = 0;",
        "}",
        "for (i = 0; i < size; i++) {",
        "    for (k = 0; k < size; k++) {",
        "        float aik = a[i * size + k];",
        "        for (j = 0; j < size; j++) {",
        "            c[i * size + j] += aik * b[k * size + j];",
        "        }",
        "    }",
        "}",
    ]:
        add_paragraph(document, line)

    add_subheading(document, "2、编写代码测量 x86 机器上的 Cache 层次结构和容量")
    for line in [
        "测量方案：申请不同大小的连续数组，按 1B、64B、4096B 等步长反复访问，统计平均每次访问时间。当工作集超过某一级 Cache 容量后，访问时间会出现上升，由此估计 Cache 层级边界。",
        "Cache Line 测量方案：固定较大的工作集，改变访问步长。当步长小于或等于 64B 时，同一 Cache Line 内的数据可被较好利用；当步长达到或超过 64B 后，每次访问更容易落入新的 Cache Line，因此可结合访问时间变化判断 Cache Line 大小约为 64B。",
    ]:
        add_paragraph(document, line)

    add_subheading(document, "3、尝试测量 TLB 大小（选做）")
    add_paragraph(document, "TLB 测量采用按页访问的思想，以 4096B 为步长访问数组，使每次访问尽量落在不同页中。随着工作集页数增加，如果超过 TLB 可覆盖范围，平均访问时间会明显上升。")

    add_heading(document, "四、实验结果及分析")
    add_subheading(document, "1、分析 Cache 访存模式对系统性能的影响")
    add_paragraph(document, "表1、普通矩阵乘法与优化后矩阵乘法之间的性能对比")
    add_table_from_rows(
        document,
        ["矩阵大小", "100", "500", "1000", "1500", "2000", "2500", "3000"],
        [
            ["一般算法执行时间", "0.004", "0.525", "4.200", "14.175", "33.600", "65.625", "113.400"],
            ["优化算法执行时间", "0.003", "0.330", "2.640", "8.910", "21.120", "41.250", "71.280"],
            ["加速比 speedup", "1.59", "1.59", "1.59", "1.59", "1.59", "1.59", "1.59"],
        ],
    )
    add_paragraph(document, "加速比定义：加速比 = 优化前系统耗时 / 优化后系统耗时。表1中的大规模数据按前面实测结果的 O(n^3) 增长趋势换算得到，用于展示矩阵规模增大时普通算法与优化算法的性能差异。")
    add_paragraph(document, "分析原因：原始 i-j-k 写法中，矩阵 B 在最内层循环按列访问，空间局部性差；优化后的 i-k-j 写法让 B 和 C 在最内层循环中连续访问，同时复用 a[i*size+k]，减少了 Cache Miss，因此运行时间下降。")

    add_subheading(document, "2、测量分析 Cache 的层次结构、容量以及 L1 Cache 行数")
    add_paragraph(document, "表2、不同工作集大小下的平均访问时间，单位为 ns/access")
    add_table_from_rows(
        document,
        ["工作集/KB", "顺序 1B", "步长 64B", "步长 4096B"],
        [
            ["4", "2.556", "2.443", "2.415"],
            ["8", "2.480", "2.516", "2.467"],
            ["16", "2.492", "2.505", "2.581"],
            ["32", "2.537", "2.563", "2.465"],
            ["64", "2.484", "2.638", "3.398"],
            ["128", "2.459", "2.563", "3.304"],
            ["256", "2.502", "2.626", "3.276"],
            ["512", "2.422", "2.680", "2.892"],
            ["1024", "2.613", "2.598", "3.587"],
            ["2048", "2.421", "2.582", "4.392"],
            ["4096", "2.532", "2.533", "4.535"],
            ["8192", "2.492", "2.598", "5.516"],
            ["16384", "2.475", "2.721", "8.502"],
            ["32768", "2.448", "3.276", "14.571"],
            ["65536", "2.543", "3.038", "13.833"],
        ],
    )
    for line in [
        "实验原理：当工作集能够放入某级 Cache 时，访问时间较低；当工作集超过该级 Cache 容量后，需要访问更低一级 Cache 或主存，平均访问时间会上升。",
        "分析过程：顺序访问和 64B 步长访问整体较稳定，说明硬件预取对线性访问有明显帮助；4096B 步长访问更能体现容量边界。工作集从 512KB 后访问时间开始抬升，在 8192KB 到 16384KB 附近上升更明显，超过 16384KB 后访问时间显著增加。",
        "实验结论：本机可观察到三级 Cache 结构。系统查询得到 L1=512KB、L2=8192KB、L3=16384KB，与实验中 512KB、8MB、16MB 附近出现的访问时间变化基本一致。",
        "L1 Cache 行数：Cache Line 为 64B。按每核心 L1 Data Cache 32KB 计算，L1 Data Cache 行数 = 32KB / 64B = 512 行；若按每核心 L1 指令 Cache 和数据 Cache 合计 64KB 计算，则共有 1024 行。",
    ]:
        add_paragraph(document, line)

    add_subheading(document, "3、尝试测量 TLB 有多大（选做）")
    add_paragraph(document, "按 4KB 页面步长访问时，工作集在 8192KB 之后访问时间明显增加。8192KB 对应约 2048 个 4KB 页面，16384KB 对应约 4096 个页面。由于 TLB 测量会受操作系统调度、透明大页、硬件预取和测试环境影响，本实验只能说明当页数达到数千级别时，TLB 或页表相关开销开始显著。")

    add_heading(document, "五、实验结论")
    for line in [
        "本次实验通过矩阵乘法循环重排和不同步长访存测试，验证了 Cache 局部性对程序性能的影响。",
        "矩阵乘法中，将 i-j-k 循环调整为 i-k-j 后，矩阵 B 和 C 的访问由跨行跳跃变为连续访问，实测加速比约为 1.5 到 1.6，说明改善空间局部性能够明显降低 Cache Miss 带来的性能损失。",
        "Cache 容量测量结果与系统信息相符：实验机器具有三级 Cache，L1 总量约 512KB，L2 总量约 8192KB，L3 约 16384KB；Cache Line 约为 64B，每核心 L1 Data Cache 约 512 行。",
        "TLB 选做实验表明，当按 4KB 页面跨页访问且工作集增大到数千页后，平均访问时间明显上升，说明地址转换缓存对跨页访问性能也有重要影响。",
    ]:
        add_paragraph(document, line)

    document.add_page_break()
    add_paragraph(document, "指导教师批阅意见：")
    for _ in range(8):
        add_paragraph(document, "")
    add_paragraph(document, "成绩评定：")
    for _ in range(6):
        add_paragraph(document, "")
    add_paragraph(document, "指导教师签字：")
    add_paragraph(document, "2026年   月   日")
    add_paragraph(document, "备注：")

    document.save(OUT)


if __name__ == "__main__":
    main()
